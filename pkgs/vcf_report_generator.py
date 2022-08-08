import pandas as pd
import docx
from docx.shared import Inches
import os
from pathlib import Path
from tempfile import TemporaryDirectory

from .report_generator_utils import plot_variants_distribution, plot_sample_variants_profiles



class VCFReportGenerator:
    ''' Class for generation of report on parsed VCF file'''
    def __init__(self, vcf_parser=None, outdir=False):
        self.vcf_parser = vcf_parser
        if not os.path.exists(outdir):
            raise OSError(f'Output directory not found: {outdir}')
        else:
            self.outdir = outdir
        self.report = None

    def create_docx_report(self):
        ''' Method creates docx object and add title to the future report '''
        self.report = docx.Document()
        self.report.add_heading(f'VCF parsing results report', 0)

    def output_metadata_metrics(self):
        ''' Method generates Metadata overview section of the report '''
        self.report.add_heading(f'Overview', 1)
        self.report.add_paragraph(f'This report describes the data that was collected from the following VCF file:')
        self.report.add_paragraph(f'{self.vcf_parser.vcf_path}')
        self.report.add_paragraph(f'The parsed VCF file is of {self.vcf_parser.metadata["fileformat"]} format and includes information on {self.vcf_parser.get_samples_number()} samples and {self.vcf_parser.get_variants_number()} variants on {self.vcf_parser.get_contigs_number()} chromosome{"s"*(self.vcf_parser.get_contigs_number() - 1)}.')
        self.report.add_paragraph(f'The reference sequence that was used for variant calling is {self.vcf_parser.metadata["reference"]}.')
        self.report.add_page_break()

    def output_variants_metrics(self):
        ''' Method generates Variant summary section of the report '''
        def add_variants_summary_figures(docx_report, vcf_parser_obj, outdir_path):
            ''' 
            Function for adding variants distribution plots for each chromosome 
            into Variant summary section of the report.
            Returns the updated docx report object.
            '''
            # create temp directory inside the output directory for storing figures
            with TemporaryDirectory(dir=outdir_path) as tmpdir:
                # iterate over contigs
                for contig in vcf_parser_obj.variants:
                    # build path to variant distribution plot
                    plot_path = os.path.join(tmpdir, f'variants_distribution.{contig}.png')

                    # get data for plotting
                    variant_type_counts_per_mb = vcf_parser_obj.get_variant_type_counts_per_mb(contig)
                    substitution_type_prc_per_mb = vcf_parser_obj.get_substitution_type_prc_per_mb(contig)

                    # generate variant distribution plot
                    plot_variants_distribution(variant_type_counts_per_mb, 
                                               substitution_type_prc_per_mb,
                                               contig,
                                               plot_path)

                    # add plots to the report
                    docx_report.add_heading(f'Distribution of variants for {contig}', 3)
                    docx_report.add_picture(plot_path, width=Inches(6))

            return docx_report

        self.report.add_heading(f'Variants summary', 1)
        self.report.add_paragraph(f'The parsed VCF file contains information on {self.vcf_parser.get_variants_number()} variants on {self.vcf_parser.get_contigs_number()} chromosome{"s"*(self.vcf_parser.get_contigs_number() - 1)}.')
        self.report.add_paragraph(f'Figures below demonstrate variants distribution across each of the chromosomes in the VCF.')
        self.report = add_variants_summary_figures(self.report, self.vcf_parser, self.outdir)
        self.report.add_page_break()

    def output_samples_metrics(self):
        ''' Method generates Samples variant profiles summary section of the report '''
        def add_samples_summary_figures(docx_report, vcf_parser_obj, outdir_path):
            ''' 
            Function for adding sample genetic variants profiles plots for each chromosome 
            into Samples variant profiles summary section of the report.
            Returns the updated docx report object.
            '''
            # create temp directory inside the output directory for storing figures
            with TemporaryDirectory(dir=outdir_path) as tmpdir:
                # iterate over samples and contigs
                for sample in vcf_parser_obj.samples:
                    for contig in vcf_parser_obj.samples[sample]:
                        # build path to variant distribution plot
                        plot_path = os.path.join(tmpdir, f'sample_profile.{sample}.{contig}.png')

                        # get data for plotting
                        sample_variant_type_counts_per_mb = vcf_parser_obj.get_sample_variant_type_counts_per_mb(sample, contig)
                        sample_substitution_type_prc_per_mb = vcf_parser_obj.get_sample_substitution_type_prc_per_mb(sample, contig)

                        # generate sample genetic variants profile plot
                        plot_sample_variants_profiles(sample_variant_type_counts_per_mb, 
                                                   sample_substitution_type_prc_per_mb,
                                                   contig,
                                                   plot_path)

                        # add plots to the report
                        docx_report.add_heading(f'{sample}\'s genetic variants profile for {contig}', 3)
                        docx_report.add_picture(plot_path, width=Inches(6))

            return docx_report

        self.report.add_heading(f'Samples variant profiles summary', 1)
        self.report.add_paragraph(f'The parsed VCF file contains genotype information on {self.vcf_parser.get_samples_number()} samples.')
        self.report.add_paragraph(f'Figures below show each sample genetic variants profiles across all chromosomes in the VCF parsed.')
        self.report = add_samples_summary_figures(self.report, self.vcf_parser, self.outdir)
        self.report.add_page_break()

    def save_report(self):
        ''' Method saves the docx report to file in the specified output directory '''
        def build_report_filename(path_to_vcf, output_directory):
            basename = Path(path_to_vcf).stem
            return os.path.join(output_directory, f'{basename}.report.docx')

        # build report filename
        report_path = build_report_filename(self.vcf_parser.vcf_path, self.outdir)

        # remove file with the same name if it exists
        if os.path.exists(report_path):
            os.remove(report_path)

        # save docx report
        self.report.save(report_path)
